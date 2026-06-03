from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")


def set_run_font(run, name="Microsoft YaHei UI") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei UI"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Title", 20, "0B2545", 0, 12),
        ("Heading 1", 15, "1F4D78", 14, 7),
        ("Heading 2", 12.5, "2E74B5", 10, 5),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei UI"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_paragraph_with_inline_code(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(70, 70, 70)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)
    return paragraph


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.12)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line if line else " ")
        set_run_font(run, "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(45, 45, 45)


def add_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", stripped):
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run)
            run.font.size = Pt(9.5)
            if r_idx == 0:
                run.bold = True
    doc.add_paragraph()


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)

    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_table(doc, table_lines)
            table_lines = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            add_code_block(doc, code_lines)
            code_lines = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            continue
        flush_table()

        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("# "):
            p = add_paragraph_with_inline_code(doc, stripped[2:], "Title")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if stripped.startswith("## "):
            add_paragraph_with_inline_code(doc, stripped[3:], "Heading 1")
            continue
        if stripped.startswith("### "):
            add_paragraph_with_inline_code(doc, stripped[4:], "Heading 2")
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            p = add_paragraph_with_inline_code(doc, numbered.group(2), "List Number")
            p.paragraph_format.left_indent = Inches(0.32)
            continue

        if stripped.startswith("- "):
            p = add_paragraph_with_inline_code(doc, stripped[2:], "List Bullet")
            p.paragraph_format.left_indent = Inches(0.32)
            continue

        add_paragraph_with_inline_code(doc, stripped)

    flush_table()
    flush_code()

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("发货单邮件草稿自动生成工具使用说明书")
        set_run_font(run)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


if __name__ == "__main__":
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
